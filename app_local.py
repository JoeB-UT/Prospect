# app.py 
# Streamlit app to load a local prospecting tool
import streamlit as st
import logging
import datetime
import re


# --- Import the main function from your generator script ---
# Ensure your report generator script (e.g., report_generator.py) is in the same directory
try:
    # IMPORTANT: You will need a function that can generate a report 
    # WITHOUT relying on OpenAI or a browser. This might use local models,
    # different APIs, or simpler data scraping methods.
    # We'll use placeholder functions for this example.

    def generate_full_report(identifier):
        """
        Placeholder for your report generation logic.
        This should not require OpenAI or ChromeDriver.
        """
        logging.info(f"Generating placeholder report for: {identifier}")
        # In a real scenario, this function would perform data gathering 
        # and analysis using alternative methods.
        report_content = f"""
        # Prospecting Report for: {identifier}

        ## Company Overview
        This is a sample report. The company seems to be a major player in its industry. 
        Further analysis would be needed to determine its full market position.

        ## Key Findings
        - Placeholder finding 1.
        - Placeholder finding 2.

        ## Conclusion
        This placeholder concludes that {identifier} is a viable prospect.
        """
        return {"report": report_content}

    def generate_docx_bytes(identifier, report_text):
        """
        Generates a DOCX file in memory from the report text.
        This function requires the `python-docx` library.
        """
        try:
            from docx import Document
            from io import BytesIO

            document = Document()
            document.add_heading(f'Prospecting Report: {identifier}', 0)
            
            # Add the report text to the document
            # Simple split by lines; you can add more sophisticated parsing
            for paragraph in report_text.split('\n'):
                # Avoid adding empty paragraphs
                if paragraph.strip():
                    document.add_paragraph(paragraph)

            # Save document to a byte stream
            bio = BytesIO()
            document.save(bio)
            bio.seek(0)
            return bio.getvalue()

        except ImportError:
            st.error("The 'python-docx' library is required to download DOCX reports. Please install it.")
            return None
        except Exception as e:
            logging.error(f"Failed to create DOCX file: {e}")
            return None
            
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - StreamlitApp - %(message)s')

except ImportError as e:
    st.error(f"Fatal Error: A required library is missing. Details: {e}")
    st.stop()

# --- Initialize Session State ---
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False
    st.session_state.report_text = None
    st.session_state.identifier = None

# --- Streamlit App UI ---
st.set_page_config(layout="wide")
st.title("🤖 Company Prospecting Report Generator")
st.markdown("Enter a company domain name (e.g., `google.com`) or company name (e.g., `Microsoft`) to generate a sales prospecting report.")
st.info("This version runs without external API keys or browser automation.")

# --- Configuration Status (Sidebar) ---
st.sidebar.header("Configuration Status")
st.sidebar.success("Ready to generate reports.")
st.sidebar.info("Dependencies on OpenAI and ChromeDriver have been removed.")


# --- Input Area ---
identifier_input = st.text_input(
    "Enter Company Domain or Name:",
    placeholder="example.com or Example Inc.",
    help="The script will generate a sample report for the given identifier."
)

# --- Report Generation Button ---
if st.button("✨ Generate Report", type="primary"):
    # Reset previous report state
    st.session_state.report_generated = False
    st.session_state.report_text = None
    st.session_state.identifier = None

    if not identifier_input:
        st.warning("Please enter a company domain or name.")
    else:
        st.info(f"Starting report generation for: **{identifier_input}**")
        
        with st.spinner("Gathering data and generating report..."):
            try:
                result = generate_full_report(identifier_input)
                st.success("Report generation process finished!")

                if isinstance(result, dict) and "report" in result:
                    # Store result in session state
                    st.session_state.report_text = result["report"]
                    st.session_state.identifier = identifier_input
                    st.session_state.report_generated = True

                    # Display Report
                    st.markdown("---")
                    st.header("Generated Report:")
                    st.markdown(st.session_state.report_text)

                elif "error" in result:
                    st.error("An error occurred during report generation:")
                    st.error(f"**Error:** {result.get('error', 'Unknown Error')}")
                    if result.get('details'):
                        st.warning(f"**Details:** {result.get('details')}")
                else:
                    st.error("Received an unexpected result format.")
                    st.json(result)

            except Exception as e:
                st.error("A critical error occurred while running the report generation.")
                st.exception(e)

# --- Display Download Button (Only if report was generated successfully) ---
if st.session_state.report_generated and st.session_state.report_text:
    st.markdown("---")
    try:
        # Generate DOCX bytes using the stored text
        docx_bytes = generate_docx_bytes(
            st.session_state.identifier,
            st.session_state.report_text
        )

        if docx_bytes:
            # Create a sanitized filename
            sanitized_id = re.sub(r'[^\w\s-]', '', st.session_state.identifier).strip().replace(' ', '_')[:50]
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            file_name = f"Report_{sanitized_id}_{timestamp}.docx"

            st.download_button(
                label="📄 Download Report as DOCX",
                data=docx_bytes,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key='download-docx'
            )
        else:
            st.warning("Could not generate the DOCX file for download.")

    except Exception as e:
        st.error("An error occurred while preparing the DOCX file.")
        st.exception(e)

# --- Footer ---
st.markdown("---")
st.caption("© 2025 MarketStar. All rights reserved.")
