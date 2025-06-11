# --- Reddit Prospect Profile Generator ---
# Using LMStudio local LLM plus brave search and scraping

import os
import time
import json
import re
import logging
import pandas as pd
from urllib.parse import urlencode, quote_plus, urljoin, urlparse, quote
import urllib.request
from datetime import datetime

# --- Environment Variable Loading ---
from dotenv import load_dotenv

# --- Web Interaction & Parsing ---
import requests
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import WebDriverException, NoSuchElementException, TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager

# --- LLM Integration (OpenAI library for LM Studio) ---
import openai # Import the base library to access specific error types like APIConnectionError
from openai import OpenAI, APIError, AuthenticationError, RateLimitError

import tempfile
import shutil

# report_generator.py (Add/Ensure these imports exist)
import io
from docx import Document

# %%
# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# %%
# --- File Paths & Working Directory ---
CHROMEDRIVER_PATH = r"C:\Users\jbennett\Tools\chromedriver.exe" # Make sure this path is correct for your system

# %%
# --- Configuration ---
load_dotenv()

# LM Studio Configuration
# LM_STUDIO_BASE_URL = os.getenv("LM_STUDIO_BASE_URL") # Added default from original if .env is missing
# LM_STUDIO_API_KEY = os.getenv("LM_STUDIO_API_KEY", "lm-studio") # Added default
# LM_STUDIO_MODEL = os.getenv("LM_STUDIO_MODEL_GEM") # Added default "LM_STUDIO_MODEL_GEM12"

# Brave Search API Credentials
# BRAVE_API_KEY = os.getenv("BRAVE_SEARCH_KEY")
# BRAVE_SEARCH_API_ENDPOINT = os.getenv("BRAVE_SEARCH_ENDPOINT")

# %%
# --- Constants ---
SELENIUM_TIMEOUT = 60
WEBSITE_TEXT_LIMIT = 25000
MAX_SUBPAGES_TO_SCRAPE = 5
REQUESTS_TIMEOUT = 1160 # Increased timeout for potentially slower local LLM responses
USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.75 Safari/537.36'

MAX_GLOBENEWSWIRE_ARTICLES = 3
GLOBENEWSWIRE_BASE_URL = "https://www.globenewswire.com"
REQUEST_DELAY = 3

# Initialize LM Studio Client
lm_studio_client = None
try:
    lm_studio_client = OpenAI(
        base_url=LM_STUDIO_BASE_URL,
        api_key=LM_STUDIO_API_KEY,
        timeout=REQUESTS_TIMEOUT,
    )
    logging.info(f"LM Studio client configured to use model: '{LM_STUDIO_MODEL}' at {LM_STUDIO_BASE_URL}")
except Exception as e:
    logging.error(f"Error configuring LM Studio client: {e}. Ensure the base URL is correct and the openai library is installed.")
    lm_studio_client = None

# Check Brave Search Configuration
# USE_BRAVE_SEARCH = False
# if not BRAVE_API_KEY or not BRAVE_SEARCH_API_ENDPOINT or BRAVE_API_KEY == "YOUR_BRAVE_SEARCH_API_KEY" or BRAVE_API_KEY == "YOUR_BRAVE_API_KEY_PLACEHOLDER":
#     logging.warning("Brave Search API Key or Endpoint not found/configured correctly (e.g., placeholder value detected). Brave Search will be skipped.")
# else:
#     logging.info("Brave Search API credentials loaded.")
#     USE_BRAVE_SEARCH = True

# %%
# --- Helper Functions ---
def setup_selenium_driver():
    """
    Sets up a Selenium WebDriver using webdriver-manager to automatically
    handle the ChromeDriver.
    """
    logging.info("Setting up Selenium WebDriver with automatic driver management...")
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument(f"--user-agent={USER_AGENT}")
    chrome_options.add_argument("--log-level=3") # Suppress console noise
    chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])

    # Create a temporary directory for the user profile
    temp_profile_dir = None
    try:
        temp_profile_dir = tempfile.mkdtemp(prefix="selenium_chrome_profile_")
        chrome_options.add_argument(f"--user-data-dir={temp_profile_dir}")
    except Exception as e:
        # This is not fatal, but log the warning
        logging.warning(f"Could not create temporary directory for Chrome profile: {e}")

    try:
        # Use ChromeDriverManager to automatically install and manage the driver
        service = Service(ChromeDriverManager().install())
        
        driver = webdriver.Chrome(service=service, options=chrome_options)
        driver.set_page_load_timeout(SELENIUM_TIMEOUT)
        
        logging.info("WebDriver setup complete using automatically managed driver.")
        return driver, temp_profile_dir

    except WebDriverException as e:
        logging.error(f"Error initializing WebDriver with webdriver-manager: {e}")
        if "session not created" in str(e).lower():
            logging.error("This can happen if the installed Chrome browser version is incompatible.")
            logging.error("On Streamlit Cloud, ensure 'google-chrome-stable' is in packages.txt.")
        return None, temp_profile_dir
        
    except Exception as e:
        logging.error(f"An unexpected error occurred during WebDriver initialization: {e}")
        return None, temp_profile_dir

def get_domain_from_name(company_name):
    logging.info(f"Attempting to guess domain for: {company_name} (basic placeholder)")
    potential_domain = company_name.lower().replace(" ", "").replace(",", "").replace(".", "")
    # This is a very basic guesser, real-world usage might need a more sophisticated approach or API
    for tld in ['.com', '.org', '.io', '.co', '.net', '.ai', '.tech']: # Added a few more common TLDs
        domain_guess = potential_domain + tld
        # In a real scenario, you might try to resolve this domain.
        # For now, just returning the first guess.
        logging.debug(f"Trying guess: {domain_guess}")
        return domain_guess # Returns the first guess immediately
    return None # Should ideally loop and check, but original code returns first

def sanitize_filename(name):
    name = re.sub(r'[\\/*?:"<>|]', "", name)
    name = name.replace(" ", "_")
    name = name.replace(".", "") # This will remove dots from domains, consider if that's intended. e.g. "example.com" -> "examplecom"
    return name[:100]

# %%
# --- Brave Search Functions ---
def fetch_brave_search_results(search_query: str, count: int = 1, extra_params: dict = None) -> dict:
    if not USE_BRAVE_SEARCH:
        return {
            "status": "error",
            "message": "Brave Search is not configured or disabled.",
            "results": []
        }
    
    if not BRAVE_API_KEY or BRAVE_API_KEY == "YOUR_BRAVE_SEARCH_API_KEY" or BRAVE_API_KEY == "YOUR_BRAVE_API_KEY_PLACEHOLDER" or not BRAVE_SEARCH_API_ENDPOINT:
        logging.error("Brave Search API key or endpoint not configured or is a placeholder.")
        return {
            "status": "error",
            "message": "Brave Search API key or endpoint not configured or is a placeholder.",
            "results": []
        }
    try:
        params = {"q": search_query, "count": count}
        if extra_params:
            params.update(extra_params)

        headers = {
            "Accept": "application/json",
            "X-Subscription-Token": BRAVE_API_KEY,
            "User-Agent": USER_AGENT
        }

        url = f"{BRAVE_SEARCH_API_ENDPOINT}?{urlencode(params)}"
        logging.info(f"Querying Brave Search API: {url}")
        req = urllib.request.Request(url, headers=headers)

        with urllib.request.urlopen(req, timeout=REQUESTS_TIMEOUT) as response:
            if response.status == 200:
                data = json.loads(response.read().decode())
                
                # Handle case where API explicitly returns no results
                if 'mixed' in data and data.get('mixed', {}).get('type') == 'no_results':
                    logging.info(f"Brave API explicitly indicated no results for query: '{search_query}'")
                    return {"status": "success", "message": f"No search results found for '{search_query}'.", "results": []}
                
                results_list = []
                # --- MODIFIED EXTRACTION LOGIC: Better handling of various Brave API response structures ---
                if data.get("news") and isinstance(data["news"].get("results"), list):
                    results_list = data["news"]["results"]
                    logging.info("Extracted results from data['news']['results']")
                elif data.get("web") and isinstance(data["web"].get("results"), list):
                    results_list = data["web"]["results"]
                    logging.info("Extracted results from data['web']['results']")
                elif data.get("discussions") and isinstance(data["discussions"].get("results"), list):
                    results_list = data["discussions"]["results"]
                    logging.info("Extracted results from data['discussions']['results']")
                elif isinstance(data.get("results"), list): 
                    results_list = data["results"]
                    logging.info("Extracted results from top-level data['results']")
                elif isinstance(data.get("hits"), list): 
                    results_list = data["hits"]
                    logging.info("Extracted results from top-level data['hits']")
                elif data.get("mixed"): 
                    logging.info("Attempting to extract results from data['mixed'] structure.")
                    mixed_content = data["mixed"]

                    # Helper function to process a potential container of results (list or dict)
                    # and extract individual result items in a standardized format.
                    def process_result_container(container_data):
                        processed_items = []
                        if not container_data:
                            return processed_items

                        # If the container itself is a list of result items
                        if isinstance(container_data, list):
                            for item in container_data:
                                if not isinstance(item, dict):
                                    logging.warning(f"Skipping non-dict item in a result list: {type(item)}")
                                    continue
                                
                                item_type = item.get("type")
                                actual_result_data = None

                                if item_type and item.get(item_type) and isinstance(item.get(item_type), dict):
                                    actual_result_data = item.get(item_type)
                                elif item_type and item.get(f"{item_type}_result") and isinstance(item.get(f"{item_type}_result"), dict):
                                    actual_result_data = item.get(f"{item_type}_result")
                                elif "title" in item and "url" in item: # Flat item
                                    actual_result_data = item
                                
                                if actual_result_data:
                                    processed_items.append(actual_result_data)
                                else:
                                    logging.warning(f"Could not extract usable data from item (type: {item_type}) in mixed sub-container. Keys: {list(item.keys())}")
                        # If the container is a dictionary, look for a 'results' list inside it
                        elif isinstance(container_data, dict):
                            if isinstance(container_data.get("results"), list):
                                # Assuming items in container_data["results"] are already in the desired flat format
                                processed_items.extend(container_data.get("results"))
                            # If the dict container itself is a single result item (less common for 'main'/'top'/'side')
                            elif "title" in container_data and "url" in container_data:
                                processed_items.append(container_data)
                            else:
                                logging.warning(f"Mixed sub-container is a dict but no 'results' list or recognized flat structure found. Keys: {list(container_data.keys())}")
                        return processed_items

                    # Check if this is a "no results" response
                    if isinstance(mixed_content, dict) and mixed_content.get('type') == 'no_results':
                        logging.info(f"Brave API indicated no results found for '{search_query}'")
                        return {"status": "success", "message": f"No search results found for '{search_query}'.", "results": []}
                    
                    if isinstance(mixed_content, list): 
                        logging.info("data['mixed'] is a list. Processing its items.")
                        results_list.extend(process_result_container(mixed_content))
                        if not results_list: # if process_result_container returned empty
                             logging.warning("data['mixed'] was a list, but no usable items were extracted after processing.")

                    elif isinstance(mixed_content, dict): 
                        logging.info(f"data['mixed'] is a dictionary. Processing keys: {list(mixed_content.keys())}")
                        
                        # Check for 'no_results' type
                        if mixed_content.get('type') == 'no_results':
                            logging.info(f"Brave API returned 'no_results' type for query: '{search_query}'")
                            return {"status": "success", "message": f"No search results found for '{search_query}'.", "results": []}
                        
                        # Check for web/news/discussions results directly in mixed content
                        if mixed_content.get("web") and isinstance(mixed_content["web"].get("results"), list):
                            results_list.extend(mixed_content["web"]["results"])
                            logging.info(f"Found {len(mixed_content['web']['results'])} results in mixed['web']['results']")
                        
                        if mixed_content.get("news") and isinstance(mixed_content["news"].get("results"), list):
                            results_list.extend(mixed_content["news"]["results"])
                            logging.info(f"Found {len(mixed_content['news']['results'])} results in mixed['news']['results']")
                        
                        if mixed_content.get("discussions") and isinstance(mixed_content["discussions"].get("results"), list):
                            results_list.extend(mixed_content["discussions"]["results"])
                            logging.info(f"Found {len(mixed_content['discussions']['results'])} results in mixed['discussions']['results']")
                        
                        # Process main, top, side sections
                        potential_containers_to_process = {}
                        for key_origin in ["main", "top", "side"]:
                            if key_origin in mixed_content:
                                content_of_key = mixed_content[key_origin]
                                potential_containers_to_process[key_origin] = content_of_key # Store for processing
                                
                                logging.info(f"  Inspecting data['mixed']['{key_origin}']:")
                                logging.info(f"    Type: {type(content_of_key).__name__}")
                                if isinstance(content_of_key, dict):
                                    logging.info(f"    Keys: {list(content_of_key.keys())}")
                                elif isinstance(content_of_key, list):
                                    logging.info(f"    Length: {len(content_of_key)}")
                                    if len(content_of_key) > 0 and isinstance(content_of_key[0], dict):
                                        logging.info(f"      Keys of first item in list: {list(content_of_key[0].keys())}")
                                else:
                                    logging.info(f"    Value: {str(content_of_key)[:200]}...") # Log a snippet for other types
                            else:
                                logging.info(f"  Key data['mixed']['{key_origin}'] not found in mixed_content.")

                        # Check if the 'searches' key exists and process it
                        if mixed_content.get('searches') and isinstance(mixed_content['searches'], list):
                            for search_item in mixed_content['searches']:
                                if isinstance(search_item, dict) and search_item.get('results') and isinstance(search_item['results'], list):
                                    results_list.extend(search_item['results'])
                                    logging.info(f"Found {len(search_item['results'])} results in mixed['searches'][i]['results']")

                        # Iterate through known structural keys like 'main', 'top', 'side'
                        for key in ["main", "top", "side"]: # Order might matter for relevance
                            if key in mixed_content:
                                logging.info(f"Processing data['mixed']['{key}']...")
                                container_data_to_process = mixed_content[key]
                                extracted_items = process_result_container(container_data_to_process)
                                if extracted_items:
                                    results_list.extend(extracted_items)
                                    logging.info(f"Extended {len(extracted_items)} results from data['mixed']['{key}'].")
                        
                        # Fallback for older dict structure if main/top/side didn't yield results or weren't primary
                        if not results_list:
                            logging.info("No results from 'main'/'top'/'side', trying direct 'results'/'web'/'news' in mixed dict as fallback.")
                            if isinstance(mixed_content.get("results"), list):
                                results_list.extend(mixed_content["results"]) # Extend, not assign
                            if mixed_content.get("web") and isinstance(mixed_content["web"].get("results"), list):
                                results_list.extend(mixed_content["web"]["results"])
                            if mixed_content.get("news") and isinstance(mixed_content["news"].get("results"), list):
                                results_list.extend(mixed_content["news"]["results"])
                            if mixed_content.get("discussions") and isinstance(mixed_content["discussions"].get("results"), list):
                                results_list.extend(mixed_content["discussions"]["results"])
                            
                            if results_list:
                                logging.info("Found results via fallback keys in data['mixed'] (dict).")
                            else:
                                logging.warning(f"Within data['mixed'] (dict), no results found in 'main'/'top'/'side' or direct fallback keys.")
                    else:
                        logging.warning(f"data['mixed'] is of unexpected type: {type(mixed_content)}")
                # --- MODIFICATION END ---
                
                # Check for empty result after all attempts
                if not results_list: 
                    # Log the structure overview for diagnosis
                    logging.info(f"No results found. API response structure: Top-level keys: {list(data.keys())}")
                    if 'mixed' in data:
                        mixed_type = type(data['mixed']).__name__
                        if isinstance(data['mixed'], dict):
                            logging.info(f"  'mixed' is a {mixed_type} with keys: {list(data['mixed'].keys())}")
                        elif isinstance(data['mixed'], list):
                            logging.info(f"  'mixed' is a {mixed_type} with length: {len(data['mixed'])}")
                        else:
                            logging.info(f"  'mixed' is a {mixed_type}")
                    
                    # Return a success with empty results, not an error
                    return {"status": "success", "message": f"No search results found or extracted for '{search_query}'.", "results": []}
                
                parsed_results = []
                if results_list:
                    logging.info(f"Attempting to parse {len(results_list)} extracted result items.")
                    for i, res in enumerate(results_list):
                        if not isinstance(res, dict):
                            logging.warning(f"Skipping non-dictionary item #{i} in results_list: {res}")
                            continue

                        title = res.get("title", "No title")
                        if title == "No title" and res.get("name"): # For items like videos
                            title = res.get("name")

                        description = res.get("description") or res.get("snippet")
                        if not description and isinstance(res.get("web"), dict) and res["web"].get("snippet"): # Common if res came from item["web"] payload
                            description = res["web"]["snippet"]
                        if not description: description = "No content available."

                        url_val = res.get("url")
                        if not url_val and isinstance(res.get("web"), dict) and res["web"].get("url"): # Common if res came from item["web"] payload
                            url_val = res["web"]["url"]
                        
                        provider_name = None
                        if isinstance(res.get("meta_url"), dict): 
                            provider_name = res.get("meta_url", {}).get("display_name")
                        elif "source" in res: 
                            provider_name = res.get("source")
                        elif res.get("profile") and isinstance(res["profile"], dict): 
                            provider_name = res["profile"].get("name")
                        
                        if not provider_name and url_val:
                            try: provider_name = urlparse(url_val).netloc
                            except Exception: pass

                        date_published_str = res.get("age")
                        
                        parsed_results.append({
                            "title": title,
                            "description": description.strip() if description else "",
                            "url": url_val,
                            "provider": provider_name or "Unknown Provider",
                            "date_published": date_published_str or "" 
                        })
                    logging.info(f"Successfully parsed {len(parsed_results)} items.")
                    return {"status": "success", "results": parsed_results}
                else: 
                    logging.info(f"No results to parse for '{search_query}'.")
                    return {"status": "success", "message": f"No search results found or extracted for '{search_query}'.", "results": []}
            else: 
                error_message_content = "Unknown error"
                try: error_message_content = response.read().decode()
                except Exception: pass 
                logging.error(f"Brave API request failed with status {response.status}: {error_message_content}")
                return {"status": "error", "message": f"Brave API request failed with status {response.status}: {error_message_content}", "results": []}

    except json.JSONDecodeError as e: 
        logging.error(f"JSON decoding failed for Brave API response: {str(e)}")
        return {"status": "error", "message": f"JSON decoding failed: {str(e)}", "results": []}
    except urllib.error.URLError as e: 
        logging.error(f"URL Error with Brave API ({BRAVE_SEARCH_API_ENDPOINT}): {str(e)}")
        return {"status": "error", "message": f"URL Error reaching Brave API: {str(e)}", "results": []}
    except Exception as e: 
        logging.error(f"An unexpected error occurred with Brave API: {str(e)}", exc_info=True)
        return {"status": "error", "message": f"An unexpected error occurred: {str(e)}", "results": []}
            
def search_brave_news(company_name):
    if not USE_BRAVE_SEARCH:
        return "[Brave Search skipped for news: Configuration missing or disabled]"

    logging.info(f"Searching Brave News for: {company_name}")
    news_snippets_str_list = []
    max_snippets = 7 
    query = f"{company_name} news"
    brave_results_data = fetch_brave_search_results(search_query=query, count=max_snippets, extra_params={"country":"US", "search_lang": "en"})

    if brave_results_data["status"] == "success" and brave_results_data["results"]:
        for article in brave_results_data["results"]:
            title = article.get('title', 'No Title')
            description = article.get('description', 'No Description')
            url = article.get('url', '')
            provider = article.get('provider', 'Unknown Provider')
            date_info = article.get('date_published', '')

            formatted_date = date_info if date_info else 'Date N/A'
            snippet = f"Title: {title}\n  Source: {provider} ({formatted_date})\n  Description: {description}\n  URL: {url}\n---\n"
            news_snippets_str_list.append(snippet)
        
        logging.info(f"  - Found {len(news_snippets_str_list)} news snippets via Brave Search.")
        return "\n".join(news_snippets_str_list)
    elif brave_results_data["status"] == "success":
        logging.info(f"  - No news results found via Brave Search for '{company_name}'.")
        return "[No relevant news results found via Brave Search]"
    else: 
        logging.error(f"  - Error during Brave News search for {company_name}: {brave_results_data['message']}")
        return f"[Brave News search failed: {brave_results_data['message']}]"

def search_brave_company_size_estimates(company_name):
    if not USE_BRAVE_SEARCH:
        return "[Brave Search for size data skipped: Configuration missing or disabled]"

    logging.info(f"Searching Brave Web Search for size data for: {company_name}")
    size_estimate_snippets = []
    max_results = 5
    query = f'"{company_name}" annual revenue employees OR "{company_name}" company size OR "{company_name}" number of employees'
    brave_results_data = fetch_brave_search_results(search_query=query, count=max_results, extra_params={"country":"US", "search_lang": "en"})

    if brave_results_data["status"] == "success" and brave_results_data["results"]:
        results_list = brave_results_data["results"]
        logging.info(f"  - Received {len(results_list)} web results from Brave Search for size query.")
        keywords_to_find = ['revenue', 'employees', '$', '€', '£', 'million', 'billion', 'headcount', 'workforce', 'staff', 'team size']
        for result in results_list:
            title = result.get('title', 'No Title')
            snippet_text = result.get('description', '') 
            url = result.get('url', '')
            provider_or_domain = result.get('provider', urlparse(url).netloc if url else 'Unknown source')

            if snippet_text and any(keyword in snippet_text.lower() for keyword in keywords_to_find):
                formatted_result = f"Title: {title}\nURL: {url} (Source: {provider_or_domain})\nSnippet: {snippet_text}\n---\n"
                size_estimate_snippets.append(formatted_result)
        
        if size_estimate_snippets:
            logging.info(f"  - Found {len(size_estimate_snippets)} potentially relevant snippets for size data via Brave Search.")
            return "\n".join(size_estimate_snippets)
        else:
            logging.info("  - No web results snippets found containing size keywords via Brave Search.")
            return "[No relevant snippets found via Brave Web Search for size data]"
    elif brave_results_data["status"] == "success":
        logging.info(f"  - No web results found via Brave Web Search for size query for '{company_name}'.")
        return "[No web results found via Brave Web Search for size data]"
    else: 
        logging.error(f"  - Error during Brave Web Search for size data for {company_name}: {brave_results_data['message']}")
        return f"[Brave Web Search for size data failed: {brave_results_data['message']}]"

# %%
# Web Site Scraping
def scrape_website_with_subpages(driver, base_url):
    logging.info(f"Scraping website: {base_url} with priority, up to {MAX_SUBPAGES_TO_SCRAPE} subpages...")
    combined_text = ""
    scraped_urls = set()
    subpages_scraped_count = 0
    wait_timeout = 15 

    if not base_url.startswith(('http://', 'https://')):
        base_url = 'https://' + base_url

    # Expanded and refined keywords
    keywords = [
        # About/Company Info (High Priority)
        'about', 'about-us', 'company', 'who-we-are', 'our-story', 'mission', 'vision', 'values', 'history', 'overview',
        # Team/Leadership (High Priority)
        'team', 'our-team', 'leadership', 'management', 'our-management', 'executives', 'board', 'directors', 'people', 'our-people', 
        'staff', 'personnel', 'staff-directory', 'faculty', 'meet-the-team', 'members', 'consultants',
        # Contact (Medium Priority)
        'contact', 'contact-us', 'contact-information', 'locations', 'offices', 'get-in-touch',
        # Products/Services (Medium Priority)
        'products', 'services', 'solutions', 'platform', 'offerings', 'expertise', 'what-we-do',
        # News/Updates (Medium Priority)
        'news', 'press', 'media', 'updates', 'blog', 'articles', 'insights', 'resources', 'publications', 'newsletter',
        # Careers (Lower Priority for this context, but can indicate growth)
        'careers', 'jobs', 'join-us', 'hiring',
        # Investor Relations (Contextual)
        'investor-relations', 'investors',
        # Customer/Client Info (Contextual)
        'clients', 'customers', 'partners', 'portfolio', 'case-studies', 'testimonials', 'reviews', 'client-stories', 'work', 'projects', 'brands',
        # Support/FAQ (Lower Priority)
        'support', 'faq', 'help',
        'governance' # from original
    ]
    keywords = [k.lower() for k in keywords]

    try:
        logging.info(f"  - Scraping homepage: {base_url}")
        driver.get(base_url)
        try:
            WebDriverWait(driver, wait_timeout).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
        except TimeoutException:
            logging.error(f"  - Error: Timed out waiting for homepage body for {base_url}. Aborting.")
            return "[Website scraping failed: Homepage body timeout]"

        homepage_text = driver.find_element(By.TAG_NAME, 'body').text
        combined_text += f"--- Homepage: {base_url} ---\n{homepage_text}\n\n"
        scraped_urls.add(base_url)

        potential_links = {}
        try:
            WebDriverWait(driver, wait_timeout).until(EC.presence_of_all_elements_located((By.TAG_NAME, 'a')))
            links = driver.find_elements(By.TAG_NAME, 'a')
            logging.info(f"  - Found {len(links)} links on homepage. Determining relevance and priority...")

            for link in links:
                href = link.get_attribute('href')
                if not href: continue

                link_text_raw = link.text
                link_text = link_text_raw.lower().strip() if link_text_raw else ""
                
                title_attr_raw = link.get_attribute('title')
                title_text = title_attr_raw.lower().strip() if title_attr_raw else ""
                
                parsed_href = urlparse(href)
                path_lower = parsed_href.path.lower() if parsed_href.path else ""
                path_segments = [seg for seg in path_lower.split('/') if seg]

                current_best_priority = float('inf')
                
                for priority_index, keyword_val in enumerate(keywords):
                    is_match = False
                    if keyword_val == path_lower.strip('/'): is_match = True
                    elif keyword_val in path_segments: is_match = True
                    elif f'/{keyword_val}' in path_lower or f'{keyword_val}/' in path_lower or f'-{keyword_val}' in path_lower or f'{keyword_val}-' in path_lower: is_match = True
                    elif keyword_val in link_text: is_match = True
                    elif keyword_val in title_text: is_match = True
                    
                    if is_match:
                        current_best_priority = min(current_best_priority, priority_index)
                        if current_best_priority == 0: break # Highest priority found
                
                if current_best_priority != float('inf'):
                    abs_url = urljoin(base_url, href)
                    # Normalize URL (remove fragment, trailing slash for comparison)
                    parsed_abs_url = urlparse(abs_url)
                    normalized_abs_url = urljoin(parsed_abs_url.scheme + "://" + parsed_abs_url.netloc, parsed_abs_url.path.rstrip('/'))

                    parsed_base_url = urlparse(base_url)
                    # Handle www and non-www consistently for domain comparison
                    base_domain = parsed_base_url.netloc.replace('www.', '')
                    link_domain = parsed_abs_url.netloc.replace('www.', '')

                    if link_domain == base_domain and \
                       parsed_abs_url.scheme in ['http', 'https'] and \
                       not abs_url.startswith('javascript:') and \
                       not parsed_abs_url.fragment and \
                       not any(abs_url.lower().endswith(ext) for ext in
                               ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.svg', '.webp',
                                '.zip', '.rar', '.tar', '.gz', '.doc', '.docx', '.xls', 
                                '.xlsx', '.ppt', '.pptx', '.mp3', '.mp4', '.avi', '.mov',
                                '.css', '.js', '.xml', '.rss', '.txt', '.json']):
                        existing_priority = potential_links.get(normalized_abs_url, float('inf'))
                        new_priority = min(existing_priority, current_best_priority)
                        if new_priority < existing_priority :
                            logging.debug(f"  - Updating priority for {normalized_abs_url} from {existing_priority} to {new_priority}")
                            potential_links[normalized_abs_url] = new_priority
                        elif normalized_abs_url not in potential_links:
                            logging.debug(f"  - Adding potential link {normalized_abs_url} with priority {new_priority}")
                            potential_links[normalized_abs_url] = new_priority
                
            relevant_urls = []
            if potential_links:
                sorted_links = sorted(potential_links.items(), key=lambda item: item[1])
                relevant_urls = [url for url, score in sorted_links]
                logging.info(f"  - Identified {len(relevant_urls)} potentially relevant subpage URLs.")
                logging.info(f"  - Top 5 prioritized URLs to check: {[url for url, score in sorted_links[:min(5, len(sorted_links))]]}")
            else:
                logging.info("  - No potentially relevant subpage URLs identified based on keywords.")

        except TimeoutException:
            logging.warning(f"  - Warning: Timed out waiting for links on homepage {base_url}.")
            relevant_urls = []
        except Exception as e:
            logging.warning(f"  - Warning: Could not reliably extract links from homepage {base_url}: {e}", exc_info=True)
            relevant_urls = []

        logging.info(f"Starting prioritized subpage scraping (limit: {MAX_SUBPAGES_TO_SCRAPE})...")
        for url in relevant_urls:
            if subpages_scraped_count >= MAX_SUBPAGES_TO_SCRAPE:
                logging.info(f"  - Reached max subpage limit ({MAX_SUBPAGES_TO_SCRAPE}). Stopping scrape.")
                break
            if url in scraped_urls:
                logging.debug(f"  - Skipping already scraped URL: {url}")
                continue
            if len(combined_text) >= WEBSITE_TEXT_LIMIT:
                logging.info("  - Reached text limit for website content. Stopping scrape.")
                break

            current_priority_val = potential_links.get(url, "N/A")
            logging.info(f"  - Scraping P{current_priority_val} subpage ({subpages_scraped_count + 1}/{MAX_SUBPAGES_TO_SCRAPE}): {url}")

            try:
                driver.get(url)
                WebDriverWait(driver, wait_timeout).until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
                
                main_content_element = None
                # Prioritize more specific content containers
                content_selectors = [
                    'article', '[role="article"]', # Semantic article
                    'main', '[role="main"]',        # Semantic main content
                    '.content', '.main-content', '.page-content', '.entry-content', '.post-content', # Common class names
                    '#content', '#main', '#page-content', # Common IDs
                    'div[class*="content"]', 'div[id*="content"]' # More generic divs
                ]
                for selector in content_selectors:
                    try:
                        main_content_element = driver.find_element(By.CSS_SELECTOR, selector)
                        logging.debug(f"    - Found content with selector: {selector}")
                        break 
                    except NoSuchElementException:
                        continue
                
                subpage_text = main_content_element.text if main_content_element else driver.find_element(By.TAG_NAME, 'body').text
                                     
                combined_text += f"\n--- Subpage (P{current_priority_val}): {url} ---\n{subpage_text}\n\n"
                scraped_urls.add(url)
                subpages_scraped_count += 1
            except TimeoutException:
                logging.warning(f"  - Warning: Timed out loading body on subpage {url}. Skipping.")
            except NoSuchElementException:
                logging.warning(f"  - Warning: Body tag (or other critical element) not found on {url}. Skipping.")
            except WebDriverException as e:
                logging.warning(f"  - Warning: WebDriverException scraping subpage {url}: {e}")
            except Exception as e:
                logging.warning(f"  - Warning: Unexpected error scraping subpage {url}: {e}", exc_info=True)

        logging.info(f"Finished scraping. Scraped homepage and {subpages_scraped_count} subpages from {base_url}.")
        return combined_text[:WEBSITE_TEXT_LIMIT]

    except WebDriverException as e:
        logging.error(f"Error during Selenium operation for {base_url}: {e}", exc_info=True)
        return "[Website scraping failed due to WebDriverException]"
    except Exception as e:
        logging.error(f"Unexpected error scraping website {base_url}: {e}", exc_info=True)
        return "[Website scraping failed due to unexpected error]"

# %%
# --- News Scraping Functions (GlobeNewswire) ---
def get_globenewswire_article_content(session, article_url):
    logging.info(f"  Fetching GlobeNewswire article content from: {article_url}")
    headers = {'User-Agent': USER_AGENT}
    try:
        time.sleep(REQUEST_DELAY)
        response = session.get(article_url, headers=headers, timeout=REQUESTS_TIMEOUT)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")

        # Try specific itemprop first
        article_content_div = soup.find("div", itemprop="articleBody")
        if article_content_div:
            elements = article_content_div.find_all(['p', 'ul', 'ol', 'li', 'h2', 'h3', 'h4', 'div'], recursive=True)
            if elements:
                content_parts = []
                for part in elements:
                    if part.name == 'div' and (part.find('p') or part.find('ul')):
                        continue
                    text_content = part.get_text(separator=" ", strip=True)
                    if text_content:
                        content_parts.append(text_content)
                
                full_content = "\n\n".join(filter(None, content_parts))
                full_content = full_content.replace(" ", " ")
                full_content = re.sub(r'\s+\n', '\n', full_content)
                full_content = re.sub(r'\n\s+', '\n', full_content)
                full_content = re.sub(r'\n{3,}', '\n\n', full_content)
                if len(full_content) > 50:
                    logging.info(f"    Successfully extracted content using itemprop selector (length: {len(full_content)}).")
                    return full_content
                else:
                    logging.warning(f"    Extracted short content using itemprop and specific elements from {article_url}. Trying direct get_text on itemprop div.")
            
            plain_text = article_content_div.get_text(separator="\n", strip=True)
            if plain_text and len(plain_text) > 50:
                logging.info(f"    Successfully extracted content using itemprop selector (fallback get_text, length: {len(plain_text)}).")
                return plain_text.replace(" ", " ")

        logging.warning(f"    Could not find 'itemprop=articleBody' div or get good content from it in {article_url}. Trying fallback class selector.")
        class_selectors = ["article-body", "main-body-container article-body", "story-content", "entry-content", "article__content"]
        for class_sel in class_selectors:
            article_content_div = soup.find("div", class_=class_sel)
            if article_content_div:
                plain_text = article_content_div.get_text(separator="\n", strip=True)
                if plain_text and len(plain_text) > 50:
                    logging.info(f"    Successfully extracted content using class selector '{class_sel}' (length: {len(plain_text)}).")
                    return plain_text.replace(" ", " ")
        
        logging.warning(f"  Could not extract meaningful article content structure from {article_url} using known selectors.")
        return None
    except requests.exceptions.Timeout:
        logging.error(f"  Timeout fetching article content from {article_url}")
        return None
    except requests.exceptions.RequestException as e:
        logging.error(f"  Request error fetching article content from {article_url}: {e}")
        return None
    except Exception as e:
        logging.error(f"  Error processing article content from {article_url}: {e}", exc_info=True)
        return None

def summarize_text_with_lm_studio(text, company_name):
    global lm_studio_client
    if lm_studio_client is None:
        logging.warning("  LM Studio client not initialized. Skipping summarization.")
        return "Summarization skipped (LM Studio client not available)."
    if not text or len(text.strip()) < 100:
        logging.warning("  Skipping summarization for short or empty content.")
        return "Content too short or empty to summarize meaningfully."

    max_input_length = 12000
    if len(text) > max_input_length:
        logging.warning(f"  Text input for summarization is too long ({len(text)} chars). Truncating to {max_input_length} chars.")
        text = text[:max_input_length] + "... [TRUNCATED FOR SUMMARIZATION]"

    prompt = f"""Please provide a concise summary (target around 150-250 words) of the following news article. 
    Focus specifically on how this news relates to or affects the company '{company_name}'. 
    Highlight key points relevant for a sales team potentially engaging with this company. These could include:
    - Positive developments: Growth, new products/services, successful funding, market expansion, key hires, positive financial results.
    - Challenges or opportunities: Problems mentioned, competitive landscape shifts, new regulations affecting them, areas where they might need solutions.
    - Leadership changes or mentions of key personnel.
    - Strategic partnerships or acquisitions.
    - Market position or sentiment.
    If the direct impact on '{company_name}' is unclear, or if the company is only mentioned peripherally, please state that briefly.
    Avoid generic statements. Extract specific, actionable insights if present.

    Article Text:
    ---
    {text}
    ---
    Concise Summary for Sales Team (focused on {company_name}):
    """
    logging.info(f"  Summarizing article text for {company_name} using LM Studio (Model: {LM_STUDIO_MODEL})...")
    try:
        messages = [
            {"role": "system", "content": "You are an AI assistant specialized in accurately and concisely summarizing business news articles, extracting key insights relevant for sales professionals targeting a specific company."},
            {"role": "user", "content": prompt},
        ]
        response = lm_studio_client.chat.completions.create(
            model=LM_STUDIO_MODEL,
            messages=messages,
            max_tokens=500,
            temperature=0.5,
        )
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            summary = response.choices[0].message.content.strip()
            logging.info(f"    Summary generated successfully by LM Studio (length: {len(summary)}).")
            return summary
        else:
            finish_reason = response.choices[0].finish_reason if response.choices and hasattr(response.choices[0], 'finish_reason') else "unknown"
            logging.warning(f"  No summary content returned from LM Studio. Finish reason: {finish_reason}")
            return f"Summarization failed (No content in LM Studio AI response. Finish reason: {finish_reason})."
    except openai.APIConnectionError as e:
        logging.error(f"  LM Studio API Connection Error during summarization: {e}. Ensure LM Studio server is running at {LM_STUDIO_BASE_URL} and model '{LM_STUDIO_MODEL}' is loaded.")
        return f"Summarization failed (LM Studio Connection Error)"
    except APIError as e_api:
        status_code = e_api.status_code if hasattr(e_api, 'status_code') else "N/A"
        error_body = str(e_api.body) if hasattr(e_api, 'body') else str(e_api)
        logging.error(f"  LM Studio API Error during summarization: Status={status_code}, Body: {error_body[:200]}...")
        return f"Summarization failed (LM Studio API Error: Status {status_code})"
    except Exception as e_gen:
        logging.error(f"  Error summarizing text with LM Studio: {e_gen}", exc_info=True)
        return f"Summarization failed (Error: {str(e_gen)[:100]})"

def scrape_globenewswire_news(session, company_name):
    encoded_company_name = quote(company_name)
    search_url = f"{GLOBENEWSWIRE_BASE_URL}/en/search/keyword/{encoded_company_name}?pageSize={MAX_GLOBENEWSWIRE_ARTICLES * 2 + 5}"
    logging.info(f"Searching GlobeNewswire for '{company_name}' using URL: {search_url}")
    headers = {'User-Agent': USER_AGENT}
    articles_data = []
    processed_urls = set()

    try:
        time.sleep(REQUEST_DELAY)
        response = session.get(search_url, headers=headers, timeout=REQUESTS_TIMEOUT)
        response.raise_for_status()
        logging.info(f"  GlobeNewswire search page request successful (Status: {response.status_code})")
        soup = BeautifulSoup(response.content, "html.parser")
    except requests.exceptions.Timeout:
        logging.error(f"Timeout fetching GlobeNewswire search results for {company_name}")
        return []
    except requests.exceptions.RequestException as e:
        logging.error(f"Request error fetching GlobeNewswire search results for {company_name}: {e}")
        return []
    except Exception as e:
        logging.error(f"Unexpected error during GlobeNewswire search request for {company_name}: {e}", exc_info=True)
        return []

    news_container_div = soup.find('div', class_='results-section')
    if not news_container_div:
        news_container_div = soup.find('div', id='news-results-tabsContent')
        if not news_container_div:
             news_container_div = soup.find('div', class_='recentNewsH')
             if not news_container_div:
                logging.warning(f"Could not find primary or alternative news container div for {company_name} on GlobeNewswire.")
                return []

    article_list_items = news_container_div.find_all("li", class_=re.compile(r"\blist-result\b|\brow\b"))
    if not article_list_items:
        logging.warning(f"Could not find news list items (li.list-result or li.row) for {company_name} on GlobeNewswire.")
        return []
    
    logging.info(f"Found {len(article_list_items)} potential GlobeNewswire article list items for {company_name}.")
    article_count = 0
    for idx, item in enumerate(article_list_items):
        if article_count >= MAX_GLOBENEWSWIRE_ARTICLES:
            logging.info(f"Reached max GlobeNewswire articles ({MAX_GLOBENEWSWIRE_ARTICLES}) for {company_name}.")
            break

        logging.debug(f"  Processing list item index: {idx}")
        date_source_div = item.find("div", class_="date-source")
        main_link_div_or_h3 = item.find(["div", "h3"], class_=re.compile(r"mainLink|post-title"))

        if not date_source_div or not main_link_div_or_h3:
            logging.debug(f"    Skipping item {idx}: Missing 'div.date-source' or 'div/h3.mainLink/post-title'.")
            continue
        
        date_span = date_source_div.find("span")
        if not date_span or not date_span.text:
            logging.warning(f"    Skipping item {idx}: Could not find date span text.")
            continue
        date_text = date_span.text.strip()
        article_date_str = "Date Parse Error"
        try:
            clean_date_text = re.sub(r'\s+(ET|EST|EDT|PT|PST|PDT|CT|CST|CDT|MT|MST|MDT|GMT|UTC)$', '', date_text, flags=re.IGNORECASE).strip()
            date_formats_to_try = ["%B %d, %Y %H:%M", "%b %d, %Y %H:%M", "%Y-%m-%d %H:%M:%S"]
            parsed_date = None
            for fmt in date_formats_to_try:
                try:
                    parsed_date = datetime.strptime(clean_date_text, fmt)
                    break
                except ValueError:
                    continue
            if parsed_date:
                article_date_str = parsed_date.strftime('%Y-%m-%d %H:%M:%S')
                logging.debug(f"      Parsed date: {article_date_str} from '{date_text}'")
            else:
                raise ValueError(f"Could not parse date '{date_text}' with known formats.")
        except ValueError as ve:
            logging.warning(f"    Skipping item {idx}: {ve}")
            article_date_str = date_text

        source_link = date_source_div.find("a", class_="sourceLink")
        article_source = source_link.text.strip() if source_link and source_link.text else "Source Not Found"
        logging.debug(f"      Found source: {article_source}")

        link_element = main_link_div_or_h3.find("a")
        if not link_element or not link_element.has_attr('href') or not link_element['href']:
            logging.warning(f"    Skipping item {idx}: Could not find valid article link href.")
            continue
        relative_url = link_element["href"]
        article_url = urljoin(GLOBENEWSWIRE_BASE_URL, relative_url) if not relative_url.startswith('http') else relative_url
        
        article_title = link_element.text.strip() if link_element.text else "Title Not Found"
        logging.debug(f"      Found article URL: {article_url}")
        logging.debug(f"      Found article Title: {article_title}")

        if article_url in processed_urls:
            logging.debug(f"    Skipping duplicate URL: {article_url}")
            continue
        processed_urls.add(article_url)

        logging.info(f"Processing GlobeNewswire article {article_count + 1}/{MAX_GLOBENEWSWIRE_ARTICLES} for {company_name}: \"{article_title[:70]}...\"")
        article_content = get_globenewswire_article_content(session, article_url)

        if article_content:
            summary = summarize_text_with_lm_studio(article_content, company_name)
            articles_data.append({
                "title": article_title,
                "date": article_date_str,
                "source": article_source,
                "url": article_url,
                "summary": summary,
                "content": article_content,
            })
            article_count += 1
            logging.info(f"  Successfully processed and summarized GlobeNewswire article {article_count} for {company_name}.")
        else:
            logging.warning(f"  Skipping GlobeNewswire article because content could not be retrieved: {article_url}")
        
        if article_count < MAX_GLOBENEWSWIRE_ARTICLES: time.sleep(0.2)

    logging.info(f"Finished GlobeNewswire processing for {company_name}. Collected {len(articles_data)} articles.")
    return articles_data

# %%
# LLM Functions for Estimates and Analysis

def get_llm_company_estimates(company_name, client_instance, base_url_for_email_guess):
    if not client_instance:
        logging.warning("Cannot get LLM estimates: LM Studio client not available.")
        return "[LLM estimation skipped: LM Studio client not configured]"

    logging.info(f"Requesting LLM estimation for {company_name} using LM Studio (Model: {LM_STUDIO_MODEL})...")
    
    email_domain_guess_source = "unknown"
    if base_url_for_email_guess:
        parsed_url = urlparse(base_url_for_email_guess)
        email_domain_for_guess = parsed_url.netloc if parsed_url.netloc else company_name.lower().replace(' ', '') + ".com (guessed)"
        email_domain_guess_source = f"derived from {base_url_for_email_guess}"
    else:
        email_domain_for_guess = company_name.lower().replace(' ', '').replace('.', '') + ".com (guessed)"
        email_domain_guess_source = "guessed from company name"

    prompt = f"""You are an AI assistant. Based on publicly available information and your general knowledge up to your last training data, please provide concise estimates for the company named '{company_name}'.
If available, consider its potential website: '{base_url_for_email_guess if base_url_for_email_guess else "not provided"}'.

Please include the following, each on a new line:
1.  **Approximate Annual Revenue Range:** (e.g., <$1M, $1M-$10M, $10M-$50M, $50M-$250M, $250M-$1B, $1B+, or "Revenue estimate unavailable")
2.  **Estimated Number of Employees Range:** (e.g., 1-10, 11-50, 51-200, 201-500, 501-1000, 1001-5000, 5000+, or "Employee estimate unavailable")
3.  **Common Email Format Convention:** For a company like '{company_name}' with a potential email domain '{email_domain_for_guess}' (this domain was {email_domain_guess_source}), suggest the most common email naming convention. (e.g., firstname.lastname@{email_domain_for_guess}, f.lastname@{email_domain_for_guess}, firstinitiallastname@{email_domain_for_guess}, firstname@{email_domain_for_guess}). If highly speculative or unknown, state "Email format convention unknown or highly speculative".

If reliable estimates for any of these points are not readily available or are highly speculative, please clearly state that for the specific metric (e.g., "Revenue estimate unavailable due to limited public data.").
Be concise and provide only the requested information in the format above.
"""

    messages = [
        {"role": "system", "content": "You are an AI assistant providing company size estimations and typical email format conventions based on general public knowledge."},
        {"role": "user", "content": prompt}
    ]
    try:
        response = client_instance.chat.completions.create(
            model=LM_STUDIO_MODEL,
            messages=messages,
            max_tokens=300, 
            temperature=0.3,
        )
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            estimation_text = response.choices[0].message.content.strip()
            logging.info(f"LM Studio Estimation received for {company_name}.")
            return estimation_text
        else:
            finish_reason = response.choices[0].finish_reason if response.choices and hasattr(response.choices[0], 'finish_reason') else "unknown"
            logging.warning(f"LM Studio Estimation response empty for {company_name}. Finish reason: {finish_reason}")
            return f"[LLM estimation failed: No content in LM Studio response. Finish reason: {finish_reason}]"
    except openai.APIConnectionError as e:
        logging.error(f"  LM Studio API Connection Error during estimation: {e}. Ensure LM Studio server is running at {LM_STUDIO_BASE_URL} and model '{LM_STUDIO_MODEL}' is loaded.")
        return f"[LLM estimation failed: LM Studio Connection Error]"
    except APIError as e_api:
        status_code = e_api.status_code if hasattr(e_api, 'status_code') else "N/A"
        error_body = str(e_api.body) if hasattr(e_api, 'body') else str(e_api)
        logging.error(f"  LM Studio API Error during estimation: Status={status_code}, Body: {error_body[:200]}...")
        return f"[LLM estimation failed: LM Studio API Error (Status {status_code})]"
    except Exception as e_gen:
        logging.error(f"Unexpected Error during LLM estimation for {company_name} with LM Studio: {e_gen}", exc_info=True)
        return f"[LLM estimation failed: Unexpected error ({str(e_gen)[:100]})]"

def search_brave_relevant_subreddits(session, company_name, company_topic=""):
    """
    Searches Brave Web Search for potentially relevant subreddits to advertise
    a given company and attempts to identify their member counts from snippets.
    """
    if not USE_BRAVE_SEARCH:
        logging.warning("Brave Search skipped for subreddits: Configuration missing or disabled.")
        return "[Brave Search for subreddits skipped: Configuration missing or disabled]"

    logging.info(f"Searching Brave for relevant subreddits for: {company_name} (Topic: '{company_topic if company_topic else 'General'}')")

    found_subreddit_info = []
    max_results = 10

    query_parts = [
        f'site:reddit.com "{company_name}" relevant subreddits',
        f'site:reddit.com subreddits for "{company_name}" audience',
    ]
    if company_topic:
        query_parts.append(f'site:reddit.com "{company_topic}" subreddits discussion')
        query_parts.append(f'site:reddit.com best subreddits for "{company_topic}" users')
        query_parts.append(f'site:reddit.com "{company_name}" "{company_topic}" community')
    else:
        query_parts.append(f'site:reddit.com "{company_name}" related communities')
        query_parts.append(f'site:reddit.com discuss "{company_name}"')

    query = " OR ".join(filter(None, query_parts))
    if len(query) > 500:
        query = " OR ".join(filter(None, query_parts[:3]))
        logging.warning(f"Subreddit search query was too long, truncated to: {query}")

    try:
        logging.info(f"Querying Brave Search API (for subreddits) with query: {query}")
        brave_api_response = fetch_brave_search_results(search_query=query, count=max_results, extra_params={'country': 'US', 'search_lang': 'en'})

        if brave_api_response["status"] == "success" and brave_api_response["results"]:
            results_list = brave_api_response["results"]
            logging.info(f"  - Received {len(results_list)} web results from Brave for subreddit query.")

            subreddit_regex = re.compile(r'r/([a-zA-Z0-9_]+(?:/[a-zA-Z0-9_]+)?)')
            member_regex = re.compile(
                r'((?:\d{1,3}(?:,\d{3})*|\d+)(?:\.\d+)??\s*[kKmM]?\b\s*(?:members|subscribers|readers|users|followers|people\s+online|active\s+users|currently\s+viewing))',
                re.IGNORECASE
            )

            relevance_keywords = ['reddit', 'subreddit', 'r/', 'community', 'members', 'subscribers', 'forum']

            for result in results_list:
                title = result.get('title', 'No Title')
                snippet = result.get('description', '')
                url = result.get('url', '')
                provider = result.get('provider', urlparse(url).netloc if url else 'Unknown')

                if not ('reddit.com' in url.lower() or 'reddit.com' in provider.lower() or any(keyword in title.lower() for keyword in relevance_keywords) or any(keyword in snippet.lower() for keyword in relevance_keywords)):
                    logging.debug(f"    Skipping result not clearly related to Reddit: {title} ({url})")
                    continue

                potential_subreddits_found = subreddit_regex.findall(snippet) or subreddit_regex.findall(title) or subreddit_regex.findall(url)
                potential_subreddits = [f"r/{name}" for name in potential_subreddits_found]

                potential_member_counts_raw = member_regex.findall(snippet) or member_regex.findall(title)
                potential_member_counts = [match[0] if isinstance(match, tuple) else match for match in potential_member_counts_raw]

                if potential_subreddits:
                    formatted_result_parts = [
                        f"Source Title: {title}",
                        f"Source URL: {url} (Provider: {provider})",
                        f"Relevant Snippet: \"{snippet}\""
                    ]

                    unique_subreddits = sorted(list(set(potential_subreddits)))
                    formatted_result_parts.append(f"  Mentioned Subreddit(s) in snippet/title/URL: {', '.join(unique_subreddits)}")

                    if potential_member_counts:
                        unique_member_counts = sorted(list(set(m[0] if isinstance(m, tuple) else m for m in potential_member_counts)))
                        formatted_result_parts.append(f"  Potential Member Count(s) in Snippet/Title: {', '.join(unique_member_counts)}")
                        if len(unique_subreddits) == 1 and len(unique_member_counts) == 1:
                            formatted_result_parts.append(f"  Possible Association: {unique_subreddits[0]} with {unique_member_counts[0]}")
                    else:
                        formatted_result_parts.append("  Potential Member Count(s) in Snippet/Title: Not clearly identified.")
                    
                    formatted_result_parts.append("---\n")
                    found_subreddit_info.append("\n".join(formatted_result_parts))

            if found_subreddit_info:
                logging.info(f"  - Found {len(found_subreddit_info)} potentially relevant snippets with subreddit information.")
                return "\n".join(found_subreddit_info)
            else:
                logging.info("  - No web results snippets found containing identifiable subreddit names or member counts via Brave for this query.")
                return "[No relevant snippets found with subreddit information via Brave Web Search for this query]"
        
        elif brave_api_response["status"] == "success":
             logging.info("  - No web results returned by Brave Web Search for the subreddit query.")
             return "[No web results found via Brave Web Search for this subreddit query]"
        else:
            logging.error(f"  - Error during Brave Web Search for subreddits: {brave_api_response['message']}")
            return f"[Brave Web Search for subreddits failed: {brave_api_response['message']}]"

    except Exception as e:
        logging.exception(f"  - Unexpected error during Brave Web Search processing for subreddits ({company_name}): {e}")
        return "[Brave Web Search for subreddits failed: Unexpected processing error]"

# %%
# --- Find Social Media Links Function ---
def find_social_media_links(url):
    """
    Scrapes a given URL to find links to specified social media platforms.

    Args:
        url (str): The URL of the website to scrape.

    Returns:
        dict: A dictionary where keys are social media platform names
              and values are the URLs found on the page. If a platform
              is not found, its value will be None.
              Returns an error message string if the URL cannot be processed.
    """
    logging.info(f"Searching for social media links on: {url}")
    
    social_media_patterns = {
        "LinkedIn": "linkedin.com",
        "Twitter/X": ["twitter.com", "x.com"],
        "Facebook": "facebook.com",
        "Instagram": "instagram.com",
        "YouTube": "youtube.com",
        "TikTok": "tiktok.com",
        "Reddit": "reddit.com",
        "Whatsapp": ["whatsapp.com", "wa.me"]
    }

    found_links = {platform: None for platform in social_media_patterns.keys()}

    try:
        parsed_initial_url = urlparse(url)
        if not parsed_initial_url.scheme:
            url = "https://" + url

        headers = {
            "User-Agent": USER_AGENT,
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Connection": "keep-alive",
        }
        
        response = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
        response.raise_for_status()
        
        final_url_after_redirects = response.url
        soup = BeautifulSoup(response.content, 'html.parser')

        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href'].strip()
            if not href or href.startswith('#') or href.startswith('mailto:') or href.startswith('tel:') or href.startswith('javascript:void(0)'):
                continue

            try:
                full_url = urljoin(final_url_after_redirects, href)
                parsed_full_url = urlparse(full_url)
            except ValueError:
                continue

            if not parsed_full_url.netloc:
                continue

            # Normalize domain: remove 'www.' and convert to lowercase
            normalized_domain = parsed_full_url.netloc.lower()
            if normalized_domain.startswith('www.'):
                normalized_domain = normalized_domain[4:]

            for platform, patterns in social_media_patterns.items():
                if found_links[platform]:
                    continue

                current_patterns = patterns if isinstance(patterns, list) else [patterns]
                for pattern in current_patterns:
                    if pattern in normalized_domain or normalized_domain.startswith(pattern):
                        is_likely_social_profile = False
                        if pattern == normalized_domain or normalized_domain.endswith("." + pattern):
                            is_likely_social_profile = True

                        if is_likely_social_profile:
                            found_links[platform] = full_url
                            break
    
    except requests.exceptions.Timeout:
        logging.error(f"Timeout while trying to fetch URL {url}")
        return {"error": f"Timeout: Could not retrieve the URL {url} within the time limit."}
    except requests.exceptions.TooManyRedirects:
        logging.error(f"Too many redirects for URL {url}")
        return {"error": f"RedirectError: Too many redirects for URL {url}."}
    except requests.exceptions.RequestException as e:
        logging.error(f"Error fetching URL {url}: {e}")
        return {"error": f"RequestError: Could not retrieve or parse the URL: {url}. Details: {e}"}
    except Exception as e:
        logging.error(f"An unexpected error occurred for URL {url}: {e}")
        return {"error": f"UnexpectedError: An unexpected error occurred while processing {url}. Details: {e}"}

    return found_links

def get_social_media_links(session, company_name, company_domain):
    """
    Finds social media links for a company by directly scraping their website.
    This function replaces the previous search_brave_for_social_media function.

    Args:
        session: The requests session to use
        company_name: Name of the company
        company_domain: Domain of the company's website

    Returns:
        str: Information about found social media links, formatted for display
    """
    logging.info(f"Searching for social media links for: {company_name} (Website: {company_domain or 'Not provided'})")
    
    if not company_domain:
        logging.warning("No domain provided for social media link search")
        return "[Social media links search skipped: No domain provided]"
    
    # Ensure domain has proper URL format
    if not company_domain.startswith(('http://', 'https://')):
        company_domain = 'https://' + company_domain
    
    try:
        result = find_social_media_links(company_domain)
        
        # Check if there was an error
        if isinstance(result, dict) and "error" in result:
            logging.error(f"Error finding social media links: {result['error']}")
            return f"[Social media links search failed: {result['error']}]"
        
        # Format the results
        found_platforms = []
        for platform, link in result.items():
            if link:
                found_platforms.append(f"- {platform}: {link}")
        
        if found_platforms:
            logging.info(f"Found {len(found_platforms)} social media links for {company_name}")
            return "Social media links found by scanning website:\n" + "\n".join(found_platforms)
        else:
            logging.info(f"No social media links found for {company_name}")
            return "[No social media links found on company website]"
            
    except Exception as e:
        logging.error(f"Unexpected error when finding social media links: {e}", exc_info=True)
        return f"[Social media links search failed: {str(e)}]"

def analyze_with_llm(company_name, gathered_data, base_url_for_email_guess):
    global lm_studio_client
    if not lm_studio_client:
        logging.error("LM Studio client not configured. Skipping analysis.")
        return {"error": "LM Studio client not configured", "details": "Client object is None."}

    logging.info(f"Analyzing data for \"{company_name}\" with LM Studio (Model: {LM_STUDIO_MODEL})...")
    
    email_domain_for_llm_source = "unknown"
    if base_url_for_email_guess:
        parsed_url = urlparse(base_url_for_email_guess)
        email_domain_for_llm = parsed_url.netloc if parsed_url.netloc else company_name.lower().replace(' ', '') + ".com (guessed)"
        email_domain_for_llm_source = f"derived from {base_url_for_email_guess}"
    else:
        email_domain_for_llm = company_name.lower().replace(' ', '').replace('.', '') + ".com (guessed)"
        email_domain_for_llm_source = "guessed from company name"

    website_content_snippet = gathered_data.get('website_content', '[Website content not gathered or unavailable]')
    max_website_len = 15000 
    if len(website_content_snippet) > max_website_len:
        logging.warning(f"Website content for {company_name} truncated from {len(website_content_snippet)} to {max_website_len} chars for LM Studio prompt.")
        website_content_snippet = website_content_snippet[:max_website_len] + "\n... [TRUNCATED WEBSITE CONTENT]"

    brave_news_snippet = gathered_data.get('brave_news_snippets', '[Brave Search News skipped, failed, or returned no results]')
    
    brave_size_snippets = gathered_data.get(
        'brave_size_estimate_snippets',
        '[Brave Search Web for size data skipped, failed, or returned no relevant snippets]'
    )
    max_size_snippet_len = 4000
    if len(brave_size_snippets) > max_size_snippet_len: 
        logging.warning(f"Brave size snippets for {company_name} truncated to {max_size_snippet_len} chars for LM Studio prompt.")
        brave_size_snippets = brave_size_snippets[:max_size_snippet_len] + "\n... [TRUNCATED SIZE SNIPPETS]"

    brave_subreddits_data_for_prompt = gathered_data.get('brave_subreddits', '[Brave Subreddit search skipped, failed, or returned no results]')
    max_subreddit_len = 4000
    if len(brave_subreddits_data_for_prompt) > max_subreddit_len:
        logging.warning(f"Brave subreddit data for {company_name} truncated to {max_subreddit_len} chars for LM Studio prompt.")
        brave_subreddits_data_for_prompt = brave_subreddits_data_for_prompt[:max_subreddit_len] + "\n... [TRUNCATED SUBREDDIT DATA]"

    globenewswire_articles = gathered_data.get('globenewswire_articles', [])
    globenewswire_prompt_section = "[No relevant GlobeNewswire articles found or processed]"
    if globenewswire_articles:
        globenewswire_texts = []
        for article_idx, article in enumerate(globenewswire_articles):
            if article_idx >= 3:
                 globenewswire_texts.append("... [Additional GlobeNewswire articles truncated from prompt] ...")
                 break
            title = article.get('title', 'No Title')
            date = article.get('date', 'No Date')
            url = article.get('url', 'No URL')
            summary = article.get('summary', '')
            display_text = summary
            if not summary or "Summarization skipped" in summary or "Summarization failed" in summary or len(summary) < 50:
                content = article.get('content', '')
                display_text = (content[:1000] + "...") if len(content) > 1000 else content
                if not display_text.strip(): display_text = "[Content snippet unavailable or summary failed]"
                display_text = f"(Summary failed or too short, using content snippet): {display_text}"

            globenewswire_texts.append(f"Article {article_idx+1}:\nTitle: {title} ({date})\nURL: {url}\nSummary/Content Snippet:\n{display_text}\n---")
        if globenewswire_texts:
            globenewswire_prompt_section = "\n".join(globenewswire_texts)

    brave_social_media_info = gathered_data.get('brave_social_media_links', '[Social media link search not run, failed, or no results found]')
    max_social_media_len = 2000
    if len(brave_social_media_info) > max_social_media_len:
        logging.warning(f"Social media info for {company_name} truncated to {max_social_media_len} chars for LM Studio prompt.")
        brave_social_media_info = brave_social_media_info[:max_social_media_len] + "\n... [TRUNCATED SOCIAL MEDIA INFO]"

    prompt = f"""
You are an expert AI Sales Intelligence Analyst. Your task is to create a concise and actionable report for a Reddit advertising sales representative who is preparing to contact '{company_name}'.
The goal is to help the sales rep understand '{company_name}'s business, potential advertising needs, target audience, and how Reddit's advertising platform would benefit them.
Base your analysis on the provided data. DO NOT invent information or guess beyond interpreting the actual data.

For contemporary or up-to-date information, you can retrieve information from webpages by outputting the URL between <Internet> tags. To gather up-to-date information on a SUBJECT, output
<Internet>https://www.google.com/search?q=SUBJECT</Internet>. ONLY EVER output google urls in internet tags.

**Company Name:** {company_name}
**Potential Website (for context):** {base_url_for_email_guess if base_url_for_email_guess else "N/A"}
**Potential Email Domain (for contact ideas):** {email_domain_for_llm} (Note: This domain was {email_domain_for_llm_source})

**Provided Data for Analysis:**

--- Initial LLM Estimates (Revenue/Employees/Email Format) ---
{gathered_data.get('llm_estimates', '[Initial LLM estimation not provided or failed]')}
--- End Initial LLM Estimates ---

--- Website Content Snippet (Homepage & Key Subpages - Plain Text) ---
{website_content_snippet}
--- End Website Content Snippet ---

--- Brave Search News Snippets ---
{brave_news_snippet}
--- End Brave Search News Snippets ---

--- Brave Search Web Snippets (Potential Size Indicators) ---
{brave_size_snippets}
--- End Brave Search Web Snippets (Potential Size Indicators) ---

--- Brave Search Subreddit Snippets (Mentions of subreddits, potential member counts from search results) ---
{brave_subreddits_data_for_prompt}
--- End Brave Search Subreddit Snippets ---

--- Social Media Links ---
{brave_social_media_info}
--- End Social Media Links ---

--- GlobeNewswire Articles (Summaries/Snippets, max 3 articles) ---
{globenewswire_prompt_section}
--- End GlobeNewswire Articles ---

**Report Sections Required (Address each point concisely):**

1.  **{company_name} Profile:**
    * **Industry/Vertical/Primary Business Focus:** (e.g., B2B SaaS for cybersecurity, E-commerce for sustainable fashion, etc.)
    * **Business Overview:** Briefly describe their main products/services and apparent business model (e.g., subscription, direct sales, marketplace).
    * **Recent News Highlights & Key Takeaways for Sales:** Summarize 1-2 most relevant insights from the news that a sales rep could leverage (e.g., new product launch, funding, expansion, stated challenges).

2.  **Social Media Presence (Found on their website):**
    * Based on the "Website Content Snippet" AND the "Social Media Links" data, list any confirmed or strongly indicated official social media profiles for '{company_name}'.
    * For each platform (e.g., LinkedIn, Twitter/X, Facebook, Instagram, YouTube, TikTok, Reddit), provide the direct URL if found.
    * Prioritize links found via direct website scraping. Note if website scraping provided different or additional links.
    * If no specific official profiles are clearly identified from either source, state that clearly (e.g., "No official LinkedIn page clearly identified from provided data.").

3.  **{company_name} Size Estimate (Synthesized):**
    * Based on all provided data (LLM pre-estimation, Brave size snippets), provide a synthesized estimate for revenue and employee count. State if data is conflicting or scarce (e.g., "Revenue: $10M-$50M (estimated based on multiple sources). Employees: 50-200 (website mentions 'growing team', Brave snippets suggest ~100).").

4.  **Potential Decision Makers & Contact Info (Strictly from provided Website Text):**
    * From the "Website Content Snippet" ONLY, list any individuals explicitly mentioned with titles relevant to marketing, sales, or leadership (e.g., CEO, CMO, VP of Marketing, Head of Sales, Founder).
    * For each, provide Name, Title.
    * **Crucially: Only list Email or Phone if explicitly found next to or clearly associated with that person in the website text.** Do not guess emails here.
    * If no such contacts with explicit email/phone are found in the website text, state: "No specific employee contacts with explicit email/phone found in the provided website text."

5.  **Sample Outreach Strategy for Reddit Advertising:**
    * **Potential Digital Advertising Needs/Approach for Reddit:** Why would '{company_name}' benefit from advertising on Reddit? What specific goals could they achieve (e.g., reach niche tech communities, build brand awareness with Gen Z, drive leads for a new B2B tool)?
    * **Digital Agencies:** If '{company_name}' is an advertising agency, create a concise and persuasive pitch for the leadership team at {company_name}, an advertising agency. Explain why they should add Reddit to their standard client digital campaign portfolio.
            Include sections like these as appropriate: The Untapped Opportunity, Key Benefits for their Clients, Powerful Targeting Capabilities, and Action Plan next steps for {company_name} to start offering Reddit advertising.
    * **Relevant Subreddits (from Brave Subreddit Snippets):** List 2-3 promising subreddits mentioned in the "Brave Search Subreddit Snippets". If member counts were found in those snippets, include them. Example: "r/technology (10M members mentioned), r/smarthome (500k members mentioned)." If no specific subreddits were found, state "No specific subreddits identified in provided snippets; further research needed."
    * **Sample Email (to a hypothetical relevant contact like a Marketing Manager):** Draft a *very short* (3-4 sentences) introductory email. Reference a *specific insight* about '{company_name}' (from news or their website) and propose a brief chat about how Reddit advertising could help them reach a relevant audience or achieve a specific goal (e.g., "Saw your recent launch of Product X... Reddit's r/productXfans community could be a great place to engage early adopters.").
    * **Sample SMS Text (160 chars max):** Short, conversational SMS referencing a value prop. e.g., "{company_name} team - many [target_audience_type] discuss [relevant_topic] on Reddit. Could be a fit for your [product/service]. Interested in a quick overview? Reply YES or NO. [YourName] @ RedditAds"
    * **Sample Phone Call Intro (Voicemail if needed, <60s):** "Hi [Contact Name], this is [Your Name] from Reddit Advertising. I was impressed by [specific positive mention of company_name, e.g., their recent Series B funding / their innovative approach to X]. We're seeing companies like yours find great success engaging niche communities on Reddit, for instance, in subreddits like [mention 1-2 relevant subreddits from data if available, e.g., r/IndustrySpecific]. I believe we could help you [achieve a specific benefit, e.g., connect with early adopters for your new Y product]. My number is [Your Number], and email is [Your Email]. Thanks!"

6.  **Marketing & Sales Context (Inferred from Data):**
    * **Implied Current Marketing Activities:** Based on website content or news, what marketing activities do they seem to be doing already (e.g., content marketing, SEO, social media posting on X platform)?
    * **Likely Target Audience Profile:** Based on their products/services and any available data, who are their likely customers (e.g., SMBs in retail, software developers, environmentally conscious consumers)?
    * **Key Reddit Benefits for '{company_name}':** Concisely list 2-3 unique ways Reddit (its specific communities, ad formats, user intent) could be particularly beneficial for this specific company over other platforms.

7.  **Proposed Reddit Campaign Idea:**
    * **Primary Objective:** (e.g., Brand Awareness, Lead Generation, Community Engagement, App Installs)
    * **Target Subreddits:** (Reiterate specific subreddits from point 5, or suggest types if specific ones aren't clear from data, e.g., "Subreddits focused on [their industry/niche], technology adoption, [related hobbies/interests].")
    * **Ad Creative Approach/Messaging Focus:** (e.g., "Highlight unique feature X for tech-savvy users," "Run an AMA with their founder in r/entrepreneur," "Promote a free trial targeting users discussing [problem their product solves].")
    * **Call to Action:** (e.g., "Learn More," "Sign Up for Demo," "Join the Discussion," "Download Whitepaper.")

8.  **General Company Contact Information (If found in Website Text or News):**
    * **{company_name} Main Address:** (Street, City, State, Zip - if explicitly found)
    * **{company_name} General Phone Number:** (If explicitly found as a general contact number)
    * **{company_name} General Email Address:** (e.g., info@, sales@, contact@ - if explicitly found)
    * **{company_name} Official Website:** (Re-iterate the confirmed or primary website URL)

**Output Format:** Use clear headings for sections 1-8 as listed above. Be factual and stick to the provided data. If information for a specific point is missing or unclear from the provided data, explicitly state 'Insufficient data provided for this point' or 'Estimation based on limited data.' Avoid generic statements not substantiated by the input.
"""
    messages = [
        {"role": "system", "content": "You are an expert AI Sales Intelligence Analyst. Your task is to create a concise and actionable report for a Reddit advertising sales representative. Base your analysis on the provided data. DO NOT invent information or guess beyond interpreting the given data."},
        {"role": "user", "content": prompt}
    ]
    try:
        logging.info(f"  - Sending request to LM Studio (Model: {LM_STUDIO_MODEL}). Prompt length: ~{len(prompt)} chars.")
        response = lm_studio_client.chat.completions.create(
            model=LM_STUDIO_MODEL,
            messages=messages,
            max_tokens=3800,
            temperature=0.4,
        )
        logging.info("  - LM Studio analysis request complete.")
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            response_content = response.choices[0].message.content
            if hasattr(response, 'usage') and response.usage:
                logging.info(f"    LLM Token Usage (if provided by LM Studio): Prompt={response.usage.prompt_tokens}, Completion={response.usage.completion_tokens}, Total={response.usage.total_tokens}")
            else:
                logging.info(f"    LLM Completion received. Length: {len(response_content)} chars. (Token usage not reported by this endpoint)")
            return {"report": response_content.strip()}
        else:
            finish_reason = response.choices[0].finish_reason if response.choices and hasattr(response.choices[0], 'finish_reason') else "unknown"
            logging.warning(f"  - Warning: LM Studio response contained no choices or empty message content. Finish reason: {finish_reason}")
            error_detail = f"No valid content returned from LLM. Finish reason: {finish_reason}."
            if hasattr(response, 'model_dump_json'):
                error_detail += f" Full response dump: {response.model_dump_json(indent=2)[:500]}"
            return {"error": "LM Studio response empty or invalid", "details": error_detail}

    except openai.APIConnectionError as e:
        logging.error(f"  - LM Studio API Connection Error during analysis: {e}. Ensure LM Studio server is running at {LM_STUDIO_BASE_URL}, model '{LM_STUDIO_MODEL}' is loaded, and network is configured correctly.")
        return {"error": "LM Studio API Connection Error", "details": f"Failed to connect to LM Studio at {LM_STUDIO_BASE_URL}. Error: {str(e)}"}
    except APIError as e_api:
        status_code = e_api.status_code if hasattr(e_api, 'status_code') else "N/A"
        response_text = "N/A"
        error_body_str = "N/A"
        if hasattr(e_api, 'response') and e_api.response:
            try: response_text = e_api.response.text
            except: response_text = "Could not decode response text."
        if hasattr(e_api, 'body') and e_api.body:
            error_body_str = str(e_api.body)
        
        logging.error(f"  - LM Studio API Error during analysis: Status={status_code}. Response Text Hint='{response_text[:200]}...'. Body Hint='{error_body_str[:200]}...'")
        return {"error": "LM Studio API Error", "details": f"LM Studio API Error (Status: {status_code}). Check LM Studio server console for model '{LM_STUDIO_MODEL}'. Error: {str(e_api)}"}
    except Exception as e_gen:
        logging.error(f"  - Unexpected Error during LM Studio analysis for {company_name}: {e_gen}", exc_info=True)
        return {"error": f"LM Studio call failed: Unexpected error", "details": f"An unexpected error occurred: {str(e_gen)}"}

# %%
# Report Generation and Main Execution Logic
def generate_docx_bytes(identifier, report_text):
    logging.info(f"Generating DOCX byte stream for: {identifier}")
    try:
        document = Document()
        document.add_heading(f"Prospect Report: {identifier}", level=0)
        document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
        document.add_paragraph() 

        lines = report_text.strip().split('\n')
        current_paragraph = None
        
        for line in lines:
            stripped_line = line.strip()

            heading_match_numbered = re.match(r'^\s*\d+\s*\.\s*\*\*(.*?)\*\*:', stripped_line)
            heading_match_markdown_style = re.match(r'^(#+)\s+(.*)', stripped_line)
            bold_line_heading = re.match(r'^\s*\*\*(.*?):\*\*\s*$', stripped_line) or re.match(r'^\s*\*\*(.*?):\*\*', stripped_line)

            if heading_match_numbered:
                heading_text = heading_match_numbered.group(1).strip()
                document.add_heading(heading_text, level=1)
                current_paragraph = None
                continue
            elif heading_match_markdown_style:
                level = len(heading_match_markdown_style.group(1))
                heading_text = heading_match_markdown_style.group(2).strip().replace('**', '')
                doc_level = max(1, min(level, 4))
                if heading_text: document.add_heading(heading_text, level=doc_level)
                current_paragraph = None
                continue
            elif bold_line_heading:
                heading_text = bold_line_heading.group(1).strip()
                document.add_heading(heading_text, level=3)
                current_paragraph = None
                continue

            # Handle list items (simple bullet points starting with * or -)
            list_item_match = re.match(r'^\s*[-*]\s+(.*)', stripped_line)
            if list_item_match:
                item_text = list_item_match.group(1).strip()
                p = document.add_paragraph(style='ListBullet')
                # Process bolding within the list item
                sub_current_pos = 0
                for match in re.finditer(r'\*\*(.*?)\*\*', item_text):
                    start, end = match.span()
                    bold_text = match.group(1)
                    if start > sub_current_pos: p.add_run(item_text[sub_current_pos:start])
                    if bold_text: p.add_run(bold_text).bold = True
                    sub_current_pos = end
                if sub_current_pos < len(item_text): p.add_run(item_text[sub_current_pos:])
                current_paragraph = None
                continue

            # If it's not a heading or list item, treat as regular paragraph content
            if not stripped_line and current_paragraph is not None:
                current_paragraph = None 
                continue
            elif not stripped_line:
                continue

            if current_paragraph is None:
                current_paragraph = document.add_paragraph()
            else:
                current_paragraph = document.add_paragraph()

            # Process bold markdown (**text**) within the line
            current_pos = 0
            for match in re.finditer(r'\*\*(.*?)\*\*', stripped_line):
                start, end = match.span()
                bold_text = match.group(1)
                if start > current_pos: current_paragraph.add_run(stripped_line[current_pos:start])
                if bold_text: current_paragraph.add_run(bold_text).bold = True
                current_pos = end
            if current_pos < len(stripped_line): current_paragraph.add_run(stripped_line[current_pos:])
        
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        logging.info(f"Successfully created DOCX byte stream for {identifier}")
        return buffer.getvalue()
    except ImportError:
        logging.error("generate_docx_bytes failed: python-docx library not installed. Please install it via 'pip install python-docx'.")
        return None
    except Exception as e:
        logging.error(f"Error generating .docx bytes for {identifier}: {e}", exc_info=True)
        return None

def generate_full_report(identifier: str):
    global lm_studio_client
    logging.info(f"\n--- Starting report generation for: {identifier} ---")
    driver = None
    session = None
    temp_profile_dir_selenium = None

    if lm_studio_client is None:
        logging.error("LM Studio client is not initialized. Cannot proceed with LLM-dependent tasks.")
        return {"error": "LM Studio Client Not Initialized", "details": "Failed to configure the LM Studio client at startup. Check .env settings and LM Studio server."}

    try:
        session = requests.Session()
        session.headers.update({'User-Agent': USER_AGENT})

        company_name = None
        domain = None
        
        if not identifier or not identifier.strip():
            logging.error("Input identifier (company name or domain) cannot be empty.")
            return {"error": "Input identifier cannot be empty.", "details": ""}

        identifier = identifier.strip()
        # Basic check if identifier is likely a domain vs. a company name
        if '.' in identifier and ' ' not in identifier and len(identifier) > 3 and not identifier.endswith('.'):
            # Assumed to be a domain
            parsed_id_url = urlparse('http://' + identifier if not identifier.startswith(('http://', 'https://')) else identifier)
            domain = parsed_id_url.netloc.lower() if parsed_id_url.netloc else identifier.lower()
            
            # Try to derive a company name from the domain
            try:
                name_parts = domain.split('.')
                if len(name_parts) > 1:
                    common_tlds = ['com', 'co', 'org', 'net', 'gov', 'edu', 'io', 'ai', 'tech', 'app', 'uk', 'ca', 'de', 'fr', 'jp', 'au']
                    potential_name_part = name_parts[-2] if name_parts[-1] in common_tlds and len(name_parts) > 1 else name_parts[0]
                    if potential_name_part in ['www', 'ftp', 'mail']:
                        potential_name_part = name_parts[-2] if len(name_parts) > 2 and name_parts[-1] in common_tlds else name_parts[0]

                    company_name = potential_name_part.capitalize()
                else:
                    company_name = domain.capitalize()
            except Exception as e:
                logging.warning(f"Could not reliably derive company name from domain '{domain}': {e}. Using domain prefix.")
                company_name = domain.split('.')[0].capitalize() if '.' in domain else domain.capitalize()
            logging.info(f"Input identified as domain: {domain}, derived company name: {company_name}")
        else:
            # Assumed to be a company name
            company_name = identifier
            domain = get_domain_from_name(company_name)
            if domain:
                logging.info(f"Using domain from get_domain_from_name: {domain}")
            else:
                 logging.warning(f"Could not guess domain for {company_name}. Website scraping may be skipped or rely on search.")

        if not company_name:
            logging.error("Critical: Company name could not be determined from identifier.")
            return {"error": "Could not determine company name.", "details": "Identifier processing failed."}
        
        search_name = company_name
        base_url_for_prompts = ('https://' + domain) if domain else ""

        # Initialize raw_data with default "skipped" messages
        raw_data = {
            'website_content': "[Skipped - Domain not confirmed or scraping disabled/failed]",
            'llm_estimates': "[Skipped - LLM client issue or task skipped]",
            'brave_news_snippets': "[Skipped - Brave Search disabled or failed]",
            'brave_size_estimate_snippets': "[Skipped - Brave Search disabled or failed]",
            'brave_subreddits': "[Skipped - Brave Search disabled or failed]",
            'brave_social_media_links': "[Skipped - Social media link search not run or no results]",
            'globenewswire_articles': [],
        }

        logging.info(f"Gathering data for company: \"{search_name}\" (Domain context: {domain if domain else 'N/A'})")

        # --- LLM Pre-Estimates ---
        if lm_studio_client:
            logging.info("Fetching LLM pre-estimates using LM Studio...")
            raw_data['llm_estimates'] = get_llm_company_estimates(search_name, lm_studio_client, base_url_for_prompts)
            time.sleep(REQUEST_DELAY) 
        else:
            raw_data['llm_estimates'] = "[Skipped - LM Studio client not available]"

        # --- Website Scraping (Selenium) ---
        can_scrape_website = bool(domain) and os.path.exists(CHROMEDRIVER_PATH)
        if domain and not os.path.exists(CHROMEDRIVER_PATH):
            logging.warning(f"ChromeDriver not found at {CHROMEDRIVER_PATH}. Website scraping for '{domain}' will be skipped.")
            raw_data['website_content'] = "[Skipped - ChromeDriver not found]"
        elif not domain:
            logging.info("No confirmed domain for website scraping. It will be skipped.")
            raw_data['website_content'] = "[Skipped - Domain unknown or not confirmed]"
        
        if can_scrape_website:
            driver, temp_profile_dir_selenium = setup_selenium_driver()
            if driver and domain:
                logging.info(f"Scraping website content for: {domain}...")
                raw_data['website_content'] = scrape_website_with_subpages(driver, domain)
                time.sleep(REQUEST_DELAY)
            elif not driver:
                logging.warning(f"Proceeding without website scraping for '{domain}' due to WebDriver initialization error.")
                raw_data['website_content'] = "[Skipped - Selenium WebDriver failed to initialize]"
        
        # --- Brave Search API Calls ---
        if USE_BRAVE_SEARCH:
            logging.info("Searching Brave News API...")
            raw_data['brave_news_snippets'] = search_brave_news(search_name)
            time.sleep(REQUEST_DELAY)
            
            logging.info("Searching Brave Web Search for company size estimates...")
            raw_data['brave_size_estimate_snippets'] = search_brave_company_size_estimates(search_name)
            time.sleep(REQUEST_DELAY)
            
            company_topic_for_subreddit_search = ""
            logging.info(f"Searching Brave for relevant subreddits related to '{search_name}' (Topic: '{company_topic_for_subreddit_search if company_topic_for_subreddit_search else 'General'}')...")
            raw_data['brave_subreddits'] = search_brave_relevant_subreddits(session, search_name, company_topic=company_topic_for_subreddit_search)
            time.sleep(REQUEST_DELAY)
        else:
            logging.info("Brave Search is not configured or disabled. Skipping Brave News, Size Estimates, and Subreddit search.")

        # --- Direct Social Media Link Scraping ---
        logging.info(f"Searching for social media links on website for '{search_name}' (using domain: {domain or 'N/A'})...")
        raw_data['brave_social_media_links'] = get_social_media_links(session, search_name, domain)
        time.sleep(REQUEST_DELAY)
            
        # --- GlobeNewswire Scraping ---
        logging.info(f"Scraping GlobeNewswire for news related to '{search_name}'...")
        raw_data['globenewswire_articles'] = scrape_globenewswire_news(session, search_name)
        time.sleep(REQUEST_DELAY)

        # --- Final LLM Analysis ---
        logging.info("Starting comprehensive LLM analysis using LM Studio...")
        if lm_studio_client:
            analysis_result = analyze_with_llm(search_name, raw_data, base_url_for_prompts)
        else:
            analysis_result = {"error": "LLM Analysis Skipped", "details": "LM Studio client not available for final analysis."}
        
        logging.info(f"--- Finished processing: {identifier} ---")
        return analysis_result

    except WebDriverException as e_wd:
        logging.error(f"A WebDriver error occurred during report generation for '{identifier}': {e_wd}", exc_info=True)
        return {"error": "WebDriver Error", "details": str(e_wd)}
    except requests.exceptions.RequestException as e_req:
        logging.error(f"A network request error occurred during report generation for '{identifier}': {e_req}", exc_info=True)
        return {"error": "Network Request Error", "details": str(e_req)}
    except openai.APIError as e_openai_api:
        logging.error(f"An OpenAI API compatible error occurred (likely LM Studio) for '{identifier}': {e_openai_api}", exc_info=True)
        lm_studio_hint = f"This might be related to LM Studio (Model: {LM_STUDIO_MODEL}, URL: {LM_STUDIO_BASE_URL}). "
        return {"error": "OpenAI API Error (LM Studio)", "details": f"{lm_studio_hint}{str(e_openai_api)}"}
    except Exception as e_main:
        logging.exception(f"An unexpected critical error occurred during report generation for '{identifier}': {e_main}")
        lm_studio_hint = ""
        if "lm_studio_client" in locals() and lm_studio_client is None:
            lm_studio_hint = "LM Studio client was not initialized. "
        return {"error": "Unexpected Critical Error in Main Report Generation", "details": f"{lm_studio_hint}{str(e_main)}"}
    finally:
        if driver:
            logging.info("Closing Selenium WebDriver...")
            try: 
                driver.quit()
            except Exception as e_close_driver: 
                logging.error(f"Error closing WebDriver: {e_close_driver}")
        
        if temp_profile_dir_selenium and os.path.exists(temp_profile_dir_selenium):
            logging.info(f"Attempting to clean up Selenium temp profile dir: {temp_profile_dir_selenium}")
            try: 
                shutil.rmtree(temp_profile_dir_selenium, ignore_errors=False)
                logging.info(f"Successfully removed Selenium temp profile dir: {temp_profile_dir_selenium}")
            except Exception as e_rm_profile: 
                logging.error(f"Error removing Selenium temp profile dir {temp_profile_dir_selenium}: {e_rm_profile}. Manual cleanup might be needed.")
        
        if session:
            logging.info("Closing requests session...")
            session.close()
        logging.info(f"Resource cleanup finished for identifier: {identifier}")

# Example Usage
if __name__ == '__main__':
    # test_company = "NVIDIA" 
    # test_company = "AeroVironment"
    test_company = "reddit.com" # Test with a domain
    # test_company = "A small local bakery" # Test with a generic name

    if lm_studio_client is None:
        logging.error("LM Studio client failed to initialize at startup. Aborting test run. "
                      f"Please check LM_STUDIO_BASE_URL (currently: {LM_STUDIO_BASE_URL}), "
                      "ensure the LM Studio server is running, and the specified model is loaded.")
    else:
        logging.info(f"Starting main execution for: '{test_company}' using LM Studio.")
        report_data_result = generate_full_report(test_company)

        if isinstance(report_data_result, dict) and 'report' in report_data_result:
            final_report_text = report_data_result['report']
            print("\n\n--- GENERATED REPORT (LM STUDIO) ---")
            print(final_report_text)
            
            # Sanitize company name for filename
            safe_company_name = sanitize_filename(test_company)
            
            txt_filename = f"{safe_company_name}_report_lm_studio.txt"
            try:
                with open(txt_filename, "w", encoding="utf-8") as f:
                    f.write(f"Report for: {test_company}\nGenerated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} using LM Studio ({LM_STUDIO_MODEL})\n\n")
                    f.write(final_report_text)
                logging.info(f"Report saved to TXT: {txt_filename}")
            except IOError as e_io_txt:
                logging.error(f"Failed to save TXT report {txt_filename}: {e_io_txt}")

            # Generate and save DOCX
            docx_bytes_content = generate_docx_bytes(test_company, final_report_text)
            if docx_bytes_content:
                docx_filename = f"{safe_company_name}_report_lm_studio.docx"
                try:
                    with open(docx_filename, "wb") as f: 
                        f.write(docx_bytes_content)
                    logging.info(f"Report saved to DOCX: {docx_filename}")
                except IOError as e_io_docx:
                    logging.error(f"Failed to save DOCX report {docx_filename}: {e_io_docx}")
            else:
                logging.error("Failed to generate DOCX byte stream (generate_docx_bytes returned None).")

        elif isinstance(report_data_result, dict) and 'error' in report_data_result:
            print(f"\n\n--- ERROR DURING REPORT GENERATION ---")
            print(f"Error: {report_data_result['error']}")
            if 'details' in report_data_result and report_data_result['details']:
                print(f"Details: {report_data_result['details']}")
            
            # Specific guidance for common LM Studio issues
            if "LM Studio" in report_data_result['error'] or \
               (report_data_result.get('details') and "LM Studio" in report_data_result.get('details', '')):
                print(f"\nPlease ensure for LM Studio:\n"
                      f"1. LM Studio server is running at the configured base URL (e.g., {LM_STUDIO_BASE_URL}).\n"
                      f"2. The specified model '{LM_STUDIO_MODEL}' is downloaded AND fully loaded in LM Studio.\n"
                      f"3. The LM Studio server has network access if the script is not on the same machine.\n"
                      f"4. Check the LM Studio server console logs for any specific error messages related to the request or model loading.")
        else:
            print("\n\n--- UNKNOWN RESULT STRUCTURE ---")
            print(f"Received unexpected result: {report_data_result}")
            
    logging.info(f"Finished main execution script for: '{test_company}'")
